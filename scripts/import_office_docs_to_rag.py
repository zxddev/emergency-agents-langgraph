#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Office文档（预案Excel和案例Word）转换为JSONL格式并导入到RAG系统

作者: AI Assistant
日期: 2025-01-28
用途: 提取应急预案和历史案例，转换为RAG可索引的JSONL格式
"""

import json
import re
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

import pandas as pd
from docx import Document


def extract_disaster_type(text: str) -> str:
    """从文本中提取灾害类型"""
    disaster_map = {
        "地震": "earthquake",
        "洪涝": "flood",
        "洪水": "flood",
        "暴雨": "rainstorm",
        "滑坡": "landslide",
        "泥石流": "debris_flow",
        "塌方": "collapse",
        "内涝": "waterlogging"
    }

    for cn_name, en_name in disaster_map.items():
        if cn_name in text:
            return en_name
    return "unknown"


def extract_year_from_text(text: str) -> int:
    """从文本中提取年份"""
    match = re.search(r'20\d{2}', text)
    return int(match.group()) if match else datetime.now().year


def extract_location_from_text(text: str) -> str:
    """从文本中提取地点"""
    # 匹配省市县等地名模式
    patterns = [
        r'([\u4e00-\u9fa5]{2,}省[\u4e00-\u9fa5]{2,}[市县])',
        r'([\u4e00-\u9fa5]{2,}[市县][\u4e00-\u9fa5]{2,}[镇乡])',
        r'([\u4e00-\u9fa5]{2,}[省市县区])',
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    return "未知地点"


def process_excel_plans(excel_path: Path) -> List[Dict[str, Any]]:
    """
    处理Excel预案文件，每行作为一个独立文档

    Args:
        excel_path: Excel文件路径

    Returns:
        JSONL格式的文档列表
    """
    documents = []

    # 读取Excel的第一个sheet（包含实际数据）
    df = pd.read_excel(excel_path, sheet_name='Sheet1')

    for idx, row in df.iterrows():
        # 跳过空行
        if pd.isna(row.iloc[0]) or not str(row.iloc[0]).strip():
            continue

        plan_name = str(row['应急处置预案']).strip()
        disaster_type = extract_disaster_type(plan_name)

        # 组合所有列的内容为完整文本
        text_parts = [
            f"# {plan_name}",
            f"\n## 适用范围\n{row['适用范围']}",
            f"\n## 响应流程\n{row['响应流程']}",
            f"\n## 关键资源\n{row['关键资源']}",
            f"\n## 历史案例\n{row['历史案例']}"
        ]

        full_text = "\n".join(text_parts)

        doc = {
            "id": f"plan_{disaster_type}_{idx+1}",
            "text": full_text,
            "meta": {
                "source": "应急管理部预案库",
                "document_type": "emergency_plan",
                "disaster_type": disaster_type,
                "plan_name": plan_name,
                "extracted_date": datetime.now().isoformat()
            },
            "domain": "规范"
        }

        documents.append(doc)
        print(f"✅ 提取预案: {plan_name} ({len(full_text)} 字符)")

    return documents


def process_word_cases(docx_path: Path) -> List[Dict[str, Any]]:
    """
    处理Word案例文件，每个案例作为一个独立文档

    Word文档结构：
    - 段落包含案例标题（如"案例一：2023年京津冀特大暴雨洪水"）
    - 每个案例后跟随多个表格包含详细信息

    Args:
        docx_path: Word文件路径

    Returns:
        JSONL格式的文档列表
    """
    documents = []
    doc = Document(docx_path)

    # 提取所有案例标题
    case_titles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("案例") and "：" in text:
            case_titles.append(text)

    print(f"📊 发现 {len(case_titles)} 个案例标题")

    # 假设表格按顺序对应案例（每个案例约2-3个表格）
    # 简化处理：将所有表格内容按案例数量平均分配
    tables_per_case = len(doc.tables) // len(case_titles) if case_titles else 1

    for idx, title in enumerate(case_titles):
        # 提取元数据
        year = extract_year_from_text(title)
        location = extract_location_from_text(title)
        disaster_type = extract_disaster_type(title)

        # 提取该案例对应的表格内容
        start_table_idx = idx * tables_per_case
        end_table_idx = start_table_idx + tables_per_case

        text_parts = [f"# {title}\n"]

        # 提取表格内容
        for table_idx in range(start_table_idx, min(end_table_idx, len(doc.tables))):
            table = doc.tables[table_idx]
            text_parts.append(f"\n## 详细信息表 {table_idx - start_table_idx + 1}\n")

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # 格式化为"字段: 内容"
                if len(cells) >= 2:
                    text_parts.append(f"**{cells[0]}**: {cells[1]}")

        full_text = "\n".join(text_parts)

        doc_data = {
            "id": f"case_{year}_{disaster_type}_{idx+1}",
            "text": full_text,
            "meta": {
                "source": "应急管理部案例库",
                "document_type": "disaster_case",
                "year": year,
                "location": location,
                "disaster_type": disaster_type,
                "case_title": title,
                "extracted_date": datetime.now().isoformat()
            },
            "domain": "案例"
        }

        documents.append(doc_data)
        print(f"✅ 提取案例: {title} ({len(full_text)} 字符)")

    return documents


def save_to_jsonl(documents: List[Dict[str, Any]], output_path: Path) -> None:
    """
    保存文档列表为JSONL格式

    Args:
        documents: 文档列表
        output_path: 输出文件路径
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        for doc in documents:
            f.write(json.dumps(doc, ensure_ascii=False) + '\n')

    print(f"\n💾 已保存 {len(documents)} 个文档到: {output_path}")


def main():
    """主函数"""
    print("="*80)
    print("开始提取Office文档并转换为RAG格式")
    print("="*80 + "\n")

    # 定义文件路径
    base_dir = Path(__file__).parent.parent
    docs_dir = base_dir / "docs" / "data-management"
    output_dir = base_dir / "temp"
    output_dir.mkdir(exist_ok=True)

    excel_path = docs_dir / "预案.xlsx"
    docx_path = docs_dir / "应急案例(1).docx"

    # 检查文件是否存在
    if not excel_path.exists():
        print(f"❌ Excel文件不存在: {excel_path}")
        return
    if not docx_path.exists():
        print(f"❌ Word文件不存在: {docx_path}")
        return

    all_documents = []

    # 处理Excel预案
    print("\n📄 处理Excel预案文件...")
    print("-"*80)
    try:
        plans = process_excel_plans(excel_path)
        all_documents.extend(plans)
        print(f"\n✅ 成功提取 {len(plans)} 个预案")
    except Exception as e:
        print(f"❌ 处理Excel文件失败: {e}")
        import traceback
        traceback.print_exc()

    # 处理Word案例
    print("\n\n📄 处理Word案例文件...")
    print("-"*80)
    try:
        cases = process_word_cases(docx_path)
        all_documents.extend(cases)
        print(f"\n✅ 成功提取 {len(cases)} 个案例")
    except Exception as e:
        print(f"❌ 处理Word文件失败: {e}")
        import traceback
        traceback.print_exc()

    # 保存为JSONL
    if all_documents:
        output_path = output_dir / "emergency_docs.jsonl"
        save_to_jsonl(all_documents, output_path)

        # 打印统计信息
        print("\n" + "="*80)
        print("📊 提取统计")
        print("="*80)
        print(f"总文档数: {len(all_documents)}")
        print(f"预案数量: {len([d for d in all_documents if d['domain'] == '规范'])}")
        print(f"案例数量: {len([d for d in all_documents if d['domain'] == '案例'])}")
        print(f"\n输出文件: {output_path}")
        print(f"文件大小: {output_path.stat().st_size / 1024:.2f} KB")

        print("\n" + "="*80)
        print("✅ 转换完成！")
        print("="*80)
        print(f"\n下一步：使用以下命令将数据导入RAG系统：")
        print(f"  python -m emergency_agents.rag.cli {output_path}")
    else:
        print("\n❌ 没有提取到任何文档")


if __name__ == "__main__":
    main()
